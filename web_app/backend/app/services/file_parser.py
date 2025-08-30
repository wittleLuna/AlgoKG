"""
文件解析服务
支持多种文件格式的解析和内容提取
"""
import os
import re
import json
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
import logging
import tempfile
import shutil
import markdown
from bs4 import BeautifulSoup

# 尝试导入可选依赖
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None
    Document = None

logger = logging.getLogger(__name__)

class FileParserService:
    """文件解析服务"""
    
    def __init__(self):
        """初始化文件解析服务"""
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
        }
    
    def read_file_with_encoding(self, file_path: Path) -> str:
        """使用多种编码尝试读取文件内容"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read().strip()
                    logger.info(f"成功使用 {encoding} 编码读取文件: {file_path}")
                    return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"使用 {encoding} 编码读取文件失败: {e}")
                continue
        
        raise Exception(f"无法解码文件 {file_path}，尝试了所有编码: {encodings}")
    
    def preprocess_content(self, content: str) -> str:
        """预处理文本内容"""
        # 规范化换行符
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # 合并多余的空行
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # 规范化空格
        content = re.sub(r'[ \t]+', ' ', content)
        
        # 去除首尾空白
        content = content.strip()
        
        return content
    
    def validate_content(self, content: str, file_path: Path) -> bool:
        """验证文件内容质量"""
        if len(content) < 50:
            logger.warning(f"内容过短，跳过: {file_path.name}")
            return False
        
        # 检查是否包含算法相关内容
        algorithm_keywords = [
            '算法', '数据结构', '排序', '查找', '树', '图', '动态规划', 'dp',
            'algorithm', 'complexity', '递归', '迭代', '链表', '数组', '栈', '队列',
            '时间复杂度', '空间复杂度', '二分', '贪心', '回溯', 'dfs', 'bfs',
            'leetcode', '题目', '示例', '输入', '输出', '解法', '思路'
        ]
        
        content_lower = content.lower()
        has_algorithm_content = any(keyword in content_lower for keyword in algorithm_keywords)
        
        if not has_algorithm_content:
            logger.warning(f"未检测到算法相关内容: {file_path.name}")
            # 不直接返回 False，因为可能是其他类型的有用内容
        
        return True

    async def parse_file(self, file_path: str, file_format: str) -> Dict[str, Any]:
        """解析文件"""
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 根据文件扩展名选择解析方法
            ext = f".{file_format.lower()}" if not file_format.startswith('.') else file_format.lower()
            
            if ext not in self.supported_formats:
                # 尝试作为文本文件解析
                logger.warning(f"不支持的文件格式 {ext}，尝试作为文本文件解析")
                ext = '.txt'
            
            parser_method = self.supported_formats[ext]
            result = parser_method(file_path)
            
            # 添加通用元数据
            result['metadata'].update({
                'file_path': str(file_path_obj),
                'file_name': file_path_obj.name,
                'file_size': file_path_obj.stat().st_size,
                'parsed_at': datetime.now().isoformat(),
                'parser_version': '1.0'
            })
            
            return result
            
        except Exception as e:
            logger.error(f"文件解析失败 {file_path}: {e}")
            return {
                'content': '',
                'metadata': {
                    'error': str(e),
                    'file_path': file_path,
                    'parsed_at': datetime.now().isoformat()
                }
            }
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析纯文本文件"""
        try:
            # 使用改进的编码读取方法
            content = self.read_file_with_encoding(Path(file_path))
            
            # 预处理内容
            content = self.preprocess_content(content)
            
            # 验证内容质量
            if not self.validate_content(content, Path(file_path)):
                logger.warning(f"文件内容质量检查未通过: {file_path}")
            
            if not content:
                raise ValueError("无法解码文本文件")
            
            return {
                'content': content,
                'metadata': {
                    'format': 'text',
                    'encoding': 'auto-detected',
                    'word_count': len(content.split()),
                    'char_count': len(content),
                    'line_count': len(content.split('\n'))
                }
            }
            
        except Exception as e:
            logger.error(f"文本文件解析失败: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)}
            }
    
    def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """解析Markdown文件"""
        try:
            # 使用改进的编码读取方法
            md_content = self.read_file_with_encoding(Path(file_path))
            
            # 预处理内容
            md_content = self.preprocess_content(md_content)
            
            # 转换为HTML
            html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
            
            # 提取纯文本
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            # 提取标题
            headers = []
            for i in range(1, 7):
                headers.extend([h.get_text() for h in soup.find_all(f'h{i}')])
            
            return {
                'content': text_content,
                'metadata': {
                    'format': 'markdown',
                    'original_markdown': md_content,
                    'html_content': html_content,
                    'headers': headers,
                    'word_count': len(text_content.split()),
                    'char_count': len(text_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Markdown文件解析失败: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)}
            }
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            logger.warning("PyPDF2 不可用，无法解析PDF文件")
            return {
                'content': '',
                'metadata': {
                    'error': 'PyPDF2 库不可用，无法解析PDF文件',
                    'title': Path(file_path).stem,
                    'format': 'pdf'
                }
            }
        
        content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取元数据
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    })
                
                # 提取文本内容
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- 第 {page_num + 1} 页 ---\n"
                            content += page_text
                    except Exception as e:
                        logger.warning(f"PDF第{page_num + 1}页解析失败: {e}")
                        continue
                
                metadata.update({
                    'format': 'pdf',
                    'page_count': len(pdf_reader.pages),
                    'word_count': len(content.split()),
                    'char_count': len(content)
                })
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"PDF文件解析失败: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)}
            }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析DOCX文件"""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx 不可用，尝试使用文本方式读取")
            # 如果 docx 不可用，尝试读取为文本文件
            try:
                content = self.read_file_with_encoding(Path(file_path))
                return {
                    'content': content,
                    'metadata': {
                        'title': Path(file_path).stem,
                        'format': 'docx',
                        'note': '使用文本方式读取，可能包含格式字符'
                    }
                }
            except Exception as e:
                logger.error(f"文本方式读取DOCX失败: {e}")
                return {
                    'content': '',
                    'metadata': {'error': str(e)}
                }
        
        try:
            doc = Document(file_path)
            
            # 提取文本内容
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
            
            # 提取元数据
            metadata = {
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'word_count': len(content.split()),
                'char_count': len(content)
            }
            
            # 尝试提取文档属性
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                })
            except Exception as e:
                logger.warning(f"提取DOCX属性失败: {e}")
            
            return {
                'content': content.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"DOCX文件解析失败: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)}
            }
    
    def _parse_html(self, file_path: str) -> Dict[str, Any]:
        """解析HTML文件"""
        try:
            content = self.read_file_with_encoding(Path(file_path))
            
            # 解析HTML
            soup = BeautifulSoup(content, 'html.parser')
            
            # 提取标题
            title = ""
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text().strip()
            
            # 提取纯文本内容
            text_content = soup.get_text()
            text_content = self.preprocess_content(text_content)
            
            # 提取标题层级
            headers = []
            for i in range(1, 7):
                headers.extend([h.get_text().strip() for h in soup.find_all(f'h{i}')])
            
            return {
                'content': text_content,
                'metadata': {
                    'format': 'html',
                    'title': title,
                    'headers': headers,
                    'word_count': len(text_content.split()),
                    'char_count': len(text_content),
                    'original_html': content
                }
            }
            
        except Exception as e:
            logger.error(f"HTML文件解析失败: {e}")
            return {
                'content': '',
                'metadata': {'error': str(e)}
            }

# 全局实例
_file_parser = None

def get_file_parser() -> FileParserService:
    """获取文件解析器实例"""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParserService()
    return _file_parser
